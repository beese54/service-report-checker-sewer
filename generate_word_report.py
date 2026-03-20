"""
Generate a Word document showing pass results for all pipeline stages.
Run: python generate_word_report.py
"""
import json
import os
from pathlib import Path

import io
import tempfile

from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ─────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(r"C:\Users\allti\OneDrive\Documents\wrn_service_report_checker")
JSON_PATH  = BASE_DIR / "pipeline_output" / "full_run_2" / "pipeline_run_20260306_080308.json"
OUT_DIR    = BASE_DIR / "pipeline_output" / "full_run_2"
NO_OBS_OUT = OUT_DIR / "no_obstruction"
OUTPUT_DOC = BASE_DIR / "pipeline_output" / "full_run_2" / "WRN_Pipeline_Pass_Results.docx"

# Pass thresholds (from pipeline_config.py)
STAGE2N_MIN_INLIERS      = 10
STAGE2N_MIN_COVERAGE_PCT = 25.0
DETECTOR_CONF_ACCEPT     = 0.70
DETECTOR_CONF_REVIEW     = 0.30
STAGE3N_PASS_TIER        = "HIGH"           # washing_confidence >= 0.50
BLUR_REJECT_THRESHOLD    = 35.0
GREASE_FLAG_THRESHOLD    = 2.0
ENTROPY_CONFIRM_DELTA    = 0.3
WATER_DETECT_THRESHOLD   = 3.0

# ── Colours ────────────────────────────────────────────────────────────────────
GREEN  = RGBColor(0x0A, 0x84, 0x30)   # pass
RED    = RGBColor(0xC0, 0x20, 0x20)   # fail
AMBER  = RGBColor(0xD0, 0x80, 0x00)   # review
BLUE   = RGBColor(0x00, 0x47, 0xAB)   # section heading
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

IMG_W  = Inches(6.2)       # image width in document
THUMB_W = Inches(2.9)      # side-by-side images


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    """Set a table cell background colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_heading(doc, text, level=1, colour=BLUE):
    p    = doc.add_heading(text, level=level)
    run  = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = colour
    run.font.bold = True
    return p


def add_para(doc, text="", bold=False, italic=False, size=10, colour=None, align=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour
    if align:
        p.alignment = align
    return p


def add_pass_badge(doc, passed: bool, label=""):
    colour = GREEN if passed else RED
    tag    = "PASS" if passed else "FAIL"
    p      = doc.add_paragraph()
    r1     = p.add_run(f"  {tag}  ")
    r1.font.bold       = True
    r1.font.color.rgb  = WHITE
    r1.font.size       = Pt(11)
    r1.font.highlight_color = None
    # Colour via direct XML shading on paragraph (workaround)
    r1.font.color.rgb = colour
    if label:
        r2 = p.add_run(f"  {label}")
        r2.font.size = Pt(10)
    return p


def add_kv_table(doc, rows: list[tuple], title: str = "", pass_key: str = ""):
    """
    Add a compact key-value table.
    rows: list of (key, value, passed|None)
      passed=True -> green value, False -> red, None -> plain
    """
    if title:
        add_para(doc, title, bold=True, size=10, colour=DARK)

    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Value"
    hdr_cells[2].text = "Result"
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(c, "1A3A5C")
        c.paragraphs[0].runs[0].font.color.rgb = WHITE

    for i, (k, v, passed) in enumerate(rows):
        cells = tbl.rows[i + 1].cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        cells[2].text = "PASS" if passed is True else ("FAIL" if passed is False else "-")

        cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        r = cells[2].paragraphs[0].runs[0]
        r.font.size  = Pt(9)
        r.font.bold  = True
        if passed is True:
            r.font.color.rgb = GREEN
        elif passed is False:
            r.font.color.rgb = RED

    return tbl


MAX_IMG_PX = 1200    # max dimension in pixels before embedding
JPEG_QUALITY = 60   # JPEG quality for embedded images


def _compress_image(img_path: Path) -> io.BytesIO:
    """Resize and compress an image into an in-memory JPEG buffer."""
    img = PILImage.open(str(img_path)).convert("RGB")
    w, h = img.size
    if max(w, h) > MAX_IMG_PX:
        scale = MAX_IMG_PX / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), PILImage.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY, optimize=True)
    buf.seek(0)
    return buf


def try_add_image(doc, img_path: Path, width=IMG_W, caption=""):
    """Add image if it exists (compressed), skip silently if not."""
    if img_path and img_path.exists():
        buf = _compress_image(img_path)
        doc.add_picture(buf, width=width)
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size   = Pt(8)
            c.runs[0].font.italic = True
    else:
        add_para(doc, f"[Image not found: {img_path.name if img_path else 'N/A'}]",
                 italic=True, size=9, colour=AMBER)


def tier_colour(tier: str) -> RGBColor:
    return {
        "HIGH":   GREEN,
        "MEDIUM": AMBER,
        "LOW":    RED,
    }.get(str(tier).upper(), DARK)


# ─────────────────────────────────────────────────────────────────────────────
# Per-stage section builders
# ─────────────────────────────────────────────────────────────────────────────

def build_stage2n_section(doc, folder_name: str, s2n: dict, out_folder: Path):
    """Stage 2N — SIFT D1/D4 and U1/U4."""
    add_heading(doc, "Stage 2N — SIFT Alignment (D1/D4 & U1/U4)", level=3, colour=BLUE)
    add_para(doc,
             "Manhole Detector (YOLOv8) gates each image pair. "
             "Then SIFT keypoints are matched via FLANN + RANSAC homography. "
             f"PASS requires: RANSAC inliers >= {STAGE2N_MIN_INLIERS} "
             f"AND inlier spatial coverage >= {STAGE2N_MIN_COVERAGE_PCT}%.",
             italic=True, size=9)

    pair_stats = s2n.get("pair_stats", {})
    overall    = s2n.get("overall_pass", False)

    for prefix in ("D", "U"):
        stats = pair_stats.get(prefix)
        if not stats:
            continue

        status = stats.get("status", "")
        if status == "MISSING_IMAGES":
            add_para(doc, f"  {prefix} pair: images missing — skipped.", italic=True, size=9)
            continue

        pair_label  = stats.get("pair", f"{prefix}1/{prefix}4")
        inliers     = stats.get("ransac_inliers", 0)
        coverage    = stats.get("inlier_coverage_pct", 0.0)
        conf_before = stats.get(f"gate_conf_{prefix}1", stats.get("gate_conf_D1", 0))
        conf_after  = stats.get(f"gate_conf_{prefix}4", stats.get("gate_conf_D4", 0))
        tier_before = stats.get(f"gate_tier_{prefix}1", stats.get("gate_tier_D1", ""))
        tier_after  = stats.get(f"gate_tier_{prefix}4", stats.get("gate_tier_D4", ""))
        kp_before   = stats.get("kp_before", 0)
        kp_after    = stats.get("kp_after",  0)
        ratio_m     = stats.get("ratio_matches", 0)
        mean_diff   = stats.get("mean_diff", 0)
        pct_changed = stats.get("pct_changed", 0)
        pair_pass   = stats.get("stage2n_pass", False)

        add_para(doc, f"Pair: {pair_label}", bold=True, size=10)

        rows = [
            ("YOLOv8 confidence — Before",
             f"{conf_before:.3f} ({tier_before})",
             conf_before >= DETECTOR_CONF_ACCEPT),
            ("YOLOv8 confidence — After",
             f"{conf_after:.3f} ({tier_after})",
             conf_after >= DETECTOR_CONF_ACCEPT),
            ("Keypoints detected (before / after)",
             f"{kp_before:,} / {kp_after:,}",
             None),
            ("FLANN ratio-test matches",
             str(ratio_m),
             None),
            (f"RANSAC inliers  (threshold >= {STAGE2N_MIN_INLIERS})",
             str(inliers),
             inliers >= STAGE2N_MIN_INLIERS),
            (f"Inlier coverage  (threshold >= {STAGE2N_MIN_COVERAGE_PCT}%)",
             f"{coverage:.1f}%",
             coverage >= STAGE2N_MIN_COVERAGE_PCT),
            ("Mean pixel difference",
             f"{mean_diff:.1f}",
             None),
            ("% pixels changed",
             f"{pct_changed:.1f}%",
             None),
            ("Pair PASS",
             "YES" if pair_pass else "NO",
             pair_pass),
        ]
        add_kv_table(doc, rows, title="")
        doc.add_paragraph()

        comp_img = out_folder / f"{prefix}_composite.jpg"
        try_add_image(doc, comp_img, width=IMG_W,
                      caption=f"{pair_label} — Composite: originals + aligned comparison + methodology")
        doc.add_paragraph()

    add_para(doc,
             f"Stage 2N Overall: {'PASS' if overall else 'FAIL'}  "
             f"(at least one pair must pass)",
             bold=True, size=10,
             colour=GREEN if overall else RED)
    doc.add_paragraph()


def build_stage3n_section(doc, folder_name: str, s3n: dict, out_folder: Path):
    """Stage 3N — Washing Confidence D2/D5 & U2/U5."""
    add_heading(doc, "Stage 3N — Washing Confidence (D2/D5 & U2/U5)", level=3, colour=BLUE)
    add_para(doc,
             "Analyses intermediate image pairs (D2/D5, U2/U5) to confirm the pipe was "
             "washed. PASS requires washing_tier == 'HIGH' (confidence >= 0.50). "
             "Key signals: keypoint ratio, edge density increase, Laplacian variance "
             "increase, Shannon entropy increase, GLCM homogeneity decrease.",
             italic=True, size=9)

    pair_results = s3n.get("pair_results", {})
    overall      = s3n.get("overall_pass", False)

    for prefix in ("D2", "U2"):
        res = pair_results.get(prefix)
        if not res:
            continue

        tier    = res.get("washing_tier", "")
        conf    = res.get("washing_confidence", 0.0)
        metrics = res.get("metrics", {})
        passed  = tier == "HIGH"

        sift    = res.get("sift_stats", {})
        inliers = sift.get("ransac_inliers", 0)

        add_para(doc, f"Pair: {prefix}/{prefix.replace('2','5')}", bold=True, size=10)

        rows = [
            ("Washing tier", tier, passed),
            (f"Washing confidence  (HIGH requires >= 0.50)", f"{conf:.3f}", passed),
            ("Keypoint ratio (after/before)",
             f"{metrics.get('kp_ratio', 0):.2f}", None),
            ("Edge density increase %",
             f"{metrics.get('edge_increase_pct', 0):.1f}%", None),
            ("Laplacian variance increase %",
             f"{metrics.get('lap_increase_pct', 0):.1f}%", None),
            ("Shannon entropy increase",
             f"{metrics.get('entropy_increase', 0):.3f}", None),
            ("GLCM homogeneity — before / after",
             f"{metrics.get('homogeneity_before', 0):.3f} / "
             f"{metrics.get('homogeneity_after', 0):.3f}", None),
            ("RANSAC inliers",
             str(inliers), None),
            ("Pair PASS",
             "YES" if passed else "NO", passed),
        ]
        add_kv_table(doc, rows)
        doc.add_paragraph()

        comp_img = out_folder / f"{prefix}_composite.jpg"
        try_add_image(doc, comp_img, width=IMG_W,
                      caption=f"{prefix}/{prefix.replace('2','5')} — Composite")
        doc.add_paragraph()

    add_para(doc,
             f"Stage 3N Overall: {'PASS' if overall else 'FAIL'}",
             bold=True, size=10,
             colour=GREEN if overall else RED)
    doc.add_paragraph()


def build_stage4n_section(doc, folder_name: str, s4n: dict, out_folder: Path):
    """Stage 4N — Geometry D3/D6 & U3/U6."""
    add_heading(doc, "Stage 4N — Geometry / GF Check (D3/D6 & U3/U6)", level=3, colour=BLUE)
    add_para(doc,
             "Analyses close-up pairs (D3/D6, U3/U6). "
             f"Gate: blur >= {BLUR_REJECT_THRESHOLD}, circle found & centred. "
             f"Tracks: grease ratio (flag > {GREASE_FLAG_THRESHOLD}%), "
             f"texture entropy delta (confirm > {ENTROPY_CONFIRM_DELTA} nats), "
             f"water signature (flag > {WATER_DETECT_THRESHOLD}%).",
             italic=True, size=9)

    pair_results  = s4n.get("pair_results", {})
    folder_verdict = s4n.get("folder_verdict", "")
    overall_pass  = folder_verdict == "ACCEPTED"

    for prefix in ("D3", "U3"):
        res = pair_results.get(prefix)
        if not res:
            continue

        status  = res.get("status", "")
        score   = res.get("score", 0)
        gate    = res.get("gate", {})
        grease  = res.get("grease", {})
        texture = res.get("texture", {})
        water   = res.get("water", {})
        sift    = res.get("sift_stats", {})

        blur_ok        = gate.get("blur_score", 0) >= BLUR_REJECT_THRESHOLD
        circle_found   = gate.get("circle_found", False)
        centering_ok   = gate.get("centering_ok", False)
        grease_flagged = grease.get("flagged", False)
        entropy_ok     = texture.get("confirmed", False)
        water_detected = water.get("water_detected", False)

        add_para(doc, f"Pair: {prefix}/{prefix.replace('3','6')}", bold=True, size=10)

        rows = [
            # Gate
            (f"Blur score (Laplacian var, threshold >= {BLUR_REJECT_THRESHOLD})",
             f"{gate.get('blur_score', 0):.1f}",
             blur_ok),
            ("Manhole circle found",
             "YES" if circle_found else "NO",
             circle_found),
            ("Circle centering OK",
             "YES" if centering_ok else "NO",
             centering_ok),
            ("Gate reject reasons",
             ", ".join(gate.get("reject_reasons", [])) or "None",
             len(gate.get("reject_reasons", [])) == 0),
            # SIFT
            ("SIFT keypoints (before / after)",
             f"{sift.get('kp_before',0):,} / {sift.get('kp_after',0):,}",
             None),
            ("RANSAC inliers",
             str(sift.get("ransac_inliers", 0)),
             None),
            ("SSIM score",
             f"{sift.get('ssim_score', 0):.3f}",
             None),
            # Grease
            (f"Grease ratio  (flag > {GREASE_FLAG_THRESHOLD}%)",
             f"{grease.get('grease_pct', 0):.2f}%",
             not grease_flagged),
            # Texture
            (f"Entropy delta  (confirm > {ENTROPY_CONFIRM_DELTA} nats)",
             f"{texture.get('entropy_delta', 0):.3f}",
             entropy_ok),
            # Water
            ("Water detected",
             f"{'YES' if water_detected else 'NO'} "
             f"(conf {water.get('water_confidence', 0):.2f})",
             None),
            ("Specular highlight %",
             f"{water.get('specular_pct', 0):.2f}%",
             None),
            # Verdict
            ("Pair score (out of 4)",
             str(score),
             None),
            ("Pair status",
             status,
             status == "PASS"),
        ]
        add_kv_table(doc, rows)
        doc.add_paragraph()

        gf_img = out_folder / f"{prefix}_gf_report.jpg"
        try_add_image(doc, gf_img, width=IMG_W,
                      caption=f"{prefix}/{prefix.replace('3','6')} — GF Report")
        doc.add_paragraph()

    add_para(doc,
             f"Stage 4N Folder Verdict: {folder_verdict}",
             bold=True, size=10,
             colour=GREEN if overall_pass else AMBER)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Main document builder
# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    data    = json.loads(JSON_PATH.read_text(encoding="utf-8"))
    folders = data["folders"]
    accepted = [f for f in folders if f["status"] == "ACCEPTED"]

    print(f"Building report for {len(accepted)} ACCEPTED folders …")

    doc = Document()

    # ── Narrow margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── Cover page ─────────────────────────────────────────────────────────
    add_heading(doc, "WRN Service Report Checker", level=1, colour=DARK)
    add_heading(doc, "Pipeline Pass Results — Full Run", level=2, colour=BLUE)
    add_para(doc, f"Run date   : {data['run_date']}", size=10)
    add_para(doc, f"Input dir  : {data['input_dir']}", size=10)
    add_para(doc, f"Output dir : {data['output_dir']}", size=10)
    doc.add_paragraph()

    # Summary table
    summary = data["summary"]
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    labels = [
        ("Total folders",            summary["total"]),
        ("ACCEPTED",                 summary["accepted"]),
        ("NEEDS REVIEW",             summary["needs_review"]),
        ("OBSTRUCTION PROCESSED",    summary["obstruction_processed"]),
        ("REJECTED",                 summary["rejected"]),
        ("Accepted shown in this doc", len(accepted)),
    ]
    for i, (k, v) in enumerate(labels):
        cells = tbl.rows[i].cells
        cells[0].text = k
        cells[1].text = str(v)
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ── Legend page ────────────────────────────────────────────────────────
    add_heading(doc, "Pass / Fail Criteria Summary", level=2, colour=DARK)

    criteria = [
        ("Stage 1 — Sort / Level 1 Classification",
         "Folder classified as no_obstruction if D1+D4 OR U1+U4 present. "
         "Otherwise → obstruction path (OCR). This stage always 'passes' for ACCEPTED folders."),
        ("Stage 2N — SIFT (D1/D4 & U1/U4)",
         f"YOLOv8 gate: confidence >= {DETECTOR_CONF_ACCEPT} to auto-accept. "
         f"RANSAC inliers >= {STAGE2N_MIN_INLIERS}. "
         f"Inlier spatial coverage >= {STAGE2N_MIN_COVERAGE_PCT}%. "
         f"At least ONE pair (D or U) must pass both SIFT tests."),
        ("Stage 3N — Washing Confidence (D2/D5 & U2/U5)",
         "Uses keypoint ratio, edge density, Laplacian variance, Shannon entropy, "
         "and GLCM homogeneity to compute a washing_confidence score. "
         "PASS = tier 'HIGH' (washing_confidence >= 0.50). "
         "Both D2 and U2 pairs must pass."),
        ("Stage 4N — Geometry / GF (D3/D6 & U3/U6)",
         f"Gate: Laplacian blur >= {BLUR_REJECT_THRESHOLD}, manhole circle found & centred. "
         f"Grease flag: ratio > {GREASE_FLAG_THRESHOLD}%. "
         f"Texture confirm: entropy delta > {ENTROPY_CONFIRM_DELTA} nats. "
         f"Water: specular + blue-water > {WATER_DETECT_THRESHOLD}%. "
         f"Scored 0–4; folder verdict = ACCEPTED if no failed pairs."),
    ]

    for title, desc in criteria:
        add_para(doc, title, bold=True, size=10, colour=BLUE)
        add_para(doc, desc,  size=9, italic=True)
        doc.add_paragraph()

    doc.add_page_break()

    # ── Per-folder sections ────────────────────────────────────────────────
    for idx, folder in enumerate(accepted, 1):
        name       = folder["name"]
        detail     = folder.get("detail", {})
        out_folder = NO_OBS_OUT / name

        print(f"  [{idx}/{len(accepted)}] {name}")

        add_heading(doc,
                    f"Folder {idx} of {len(accepted)}: {name}  ——  ACCEPTED",
                    level=2, colour=GREEN)

        # Top-level report image
        summary_img = out_folder / f"{name}_report.jpg"
        gf_summary  = out_folder / f"{name}_gf_report.jpg"
        if summary_img.exists():
            try_add_image(doc, summary_img, width=IMG_W,
                          caption=f"{name} — Stage 2N/3N Summary Report")
        if gf_summary.exists():
            try_add_image(doc, gf_summary,  width=IMG_W,
                          caption=f"{name} — Stage 4N GF Summary Report")
        doc.add_paragraph()

        # Stage 2N
        if "stage2n" in detail:
            build_stage2n_section(doc, name, detail["stage2n"], out_folder)

        # Stage 3N
        if "stage3n" in detail:
            build_stage3n_section(doc, name, detail["stage3n"], out_folder)

        # Stage 4N
        if "stage4n" in detail:
            build_stage4n_section(doc, name, detail["stage4n"], out_folder)

        doc.add_page_break()

    doc.save(str(OUTPUT_DOC))
    print(f"\nSaved: {OUTPUT_DOC}")


if __name__ == "__main__":
    build_document()
